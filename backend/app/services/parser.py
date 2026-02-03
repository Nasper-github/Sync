import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from typing import List
from app.models.document import ParsedDocument, ParagraphNode, RunNode, FormatAttributes

class DocumentParser:
    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        doc = docx.Document(file_path)
        parsed_paragraphs = []
        
        # doc.paragraphs only contains body paragraphs, not tables, etc. for now.
        for p_index, paragraph in enumerate(doc.paragraphs):
            # XPath index is 1-based usually, but we can stick to 0-based or 1-based.
            # Let's use 1-based for XPath compatibility.
            p_xpath = f"/document/body/p[{p_index + 1}]"
            
            p_node = self._parse_paragraph(paragraph, p_xpath)
            parsed_paragraphs.append(p_node)
            
        return ParsedDocument(
            filename=filename,
            paragraphs=parsed_paragraphs,
            metadata={
                "core_properties": {
                    "author": doc.core_properties.author,
                    "created": str(doc.core_properties.created),
                    "modified": str(doc.core_properties.modified),
                }
            }
        )

    def _parse_paragraph(self, paragraph: Paragraph, p_xpath: str) -> ParagraphNode:
        runs = []
        for r_index, run in enumerate(paragraph.runs):
            r_xpath = f"{p_xpath}/r[{r_index + 1}]"
            runs.append(self._parse_run(run, r_xpath))
            
        return ParagraphNode(
            id=p_xpath,
            style=paragraph.style.name if paragraph.style else None,
            runs=runs
        )

    def _parse_run(self, run: Run, r_xpath: str) -> RunNode:
        formatting = FormatAttributes(
            bold=bool(run.bold),
            italic=bool(run.italic),
            underline=bool(run.underline),
            strike=bool(run.font.strike) if run.font else False
        )
        
        return RunNode(
            id=r_xpath,
            text=run.text,
            formatting=formatting
        )
